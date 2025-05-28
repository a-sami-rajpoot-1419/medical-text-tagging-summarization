#!/usr/bin/env python3

import os
import argparse
import json
import requests
import pandas as pd
from pathlib import Path
from typing import List, Dict, Tuple, Union, Optional
import pdfplumber
from docx import Document
from dotenv import load_dotenv
from tqdm import tqdm
from rouge_score import rouge_scorer
from sklearn.metrics import precision_score, recall_score, f1_score
import logging
import numpy as np
from numpy.typing import NDArray

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

class MedicalTextProcessor:
    def __init__(self, input_path: Union[str, Path], output_dir: Union[str, Path]):
        self.input_path = Path(input_path)
        self.output_dir = Path(output_dir)
        
        self.api_key = os.getenv('HUGGINGFACE_API_KEY')
        if not self.api_key:
            raise ValueError("HUGGINGFACE_API_KEY environment variable not set")
        
        self.headers = {"Authorization": f"Bearer {self.api_key}"}
        self.tagging_api_url = "https://api-inference.huggingface.co/models/distilbert-base-uncased"
        self.summarization_api_url = "https://api-inference.huggingface.co/models/facebook/bart-large-cnn"
        
        # Initialize ROUGE scorer
        self.rouge_scorer = rouge_scorer.RougeScorer(['rouge1', 'rougeL'], use_stemmer=True)

    def process_file(self) -> None:
        if not self.input_path.exists():
            raise FileNotFoundError(f"Input file not found: {self.input_path}")
        
        # Create output directory if it doesn't exist
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Process based on file type
        suffix = self.input_path.suffix.lower()
        if suffix == '.txt':
            self._process_text_file()
        elif suffix == '.docx':
            self._process_docx_file()
        elif suffix == '.pdf':
            self._process_pdf_file()
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
            
    def _process_text_file(self) -> None:
        with open(self.input_path, 'r', encoding='utf-8') as f:
            text = f.read()
        self._process_text(text)
        
    def _process_docx_file(self) -> None:
        doc = Document(str(self.input_path))  # Convert Path to string
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        self._process_text(text)
        
    def _process_pdf_file(self) -> None:
        # Add PDF processing logic here
        pass
        
    def _process_text(self, text: str) -> Dict[str, float]:
        # Your text processing logic here
        predictions = self._get_predictions(text)  # This should be defined
        
        return {
            'sentiment': float(predictions['sentiment']),
            'confidence': float(predictions['confidence']),
            'relevance': float(predictions['relevance'])
        }
        
    def _get_predictions(self, text: str) -> Dict[str, Union[float, NDArray]]:
        # Implement your prediction logic here
        return {
            'sentiment': 0.5,
            'confidence': 0.8,
            'relevance': 0.7
        }
        
    def save_results(self, results: Dict[str, float]) -> None:
        # Create output directory if it doesn't exist
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Save tags to Excel
        tags_path = str(self.output_dir / 'tags.xlsx')
        df_tags = pd.DataFrame([results])
        df_tags.to_excel(tags_path, index=False)
        
        # Save summary to Word
        summary_path = str(self.output_dir / 'summaries.docx')
        doc = Document()
        doc.add_heading('Text Analysis Summary', 0)
        for key, value in results.items():
            doc.add_paragraph(f"{key}: {value:.2f}")
        doc.save(summary_path)

    def extract_text(self, file_path: Union[str, Path]) -> str:
        """Extract text from PDF or DOCX file."""
        file_path = Path(file_path)
        
        if file_path.suffix.lower() == '.pdf':
            with pdfplumber.open(str(file_path)) as pdf:  # Convert Path to string
                text = ' '.join([page.extract_text() for page in pdf.pages])
        elif file_path.suffix.lower() == '.docx':
            doc = Document(str(file_path))  # Convert Path to string
            text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        return text

    def get_tags(self, text: str) -> List[str]:
        """Get top-3 disease tags using DistilBERT."""
        try:
            response = requests.post(
                self.tagging_api_url,
                headers=self.headers,
                json={"inputs": text}
            )
            response.raise_for_status()
            predictions = response.json()
            
            # Sort predictions by confidence and get top-3
            top_tags = sorted(predictions, key=lambda x: x['score'], reverse=True)[:3]
            return [tag['label'] for tag in top_tags]
        except Exception as e:
            logger.error(f"Error in tag prediction: {str(e)}")
            return []

    def get_summary(self, text: str) -> str:
        """Generate summary using BART."""
        try:
            response = requests.post(
                self.summarization_api_url,
                headers=self.headers,
                json={"inputs": text}
            )
            response.raise_for_status()
            summary = response.json()
            return summary[0]['summary_text']
        except Exception as e:
            logger.error(f"Error in summary generation: {str(e)}")
            return ""

    def calculate_metrics(self, predictions: List[str], references: List[str]) -> Dict[str, float]:
        """Calculate precision, recall, and F1-score."""
        # Convert to binary format for sklearn metrics
        y_true = [1 if ref in preds else 0 for preds, ref in zip(predictions, references)]
        y_pred = [1 if preds in refs else 0 for preds, refs in zip(predictions, references)]
        
        return {
            'precision': float(precision_score(y_true, y_pred)),
            'recall': float(recall_score(y_true, y_pred)),
            'f1': float(f1_score(y_true, y_pred))
        }

    def calculate_rouge_scores(self, summaries: List[str], references: List[str]) -> Dict[str, float]:
        """Calculate ROUGE-1 and ROUGE-L scores."""
        rouge1_scores = []
        rougeL_scores = []
        
        for summary, reference in zip(summaries, references):
            scores = self.rouge_scorer.score(reference, summary)
            rouge1_scores.append(scores['rouge1'].fmeasure)
            rougeL_scores.append(scores['rougeL'].fmeasure)
        
        return {
            'rouge1': float(sum(rouge1_scores) / len(rouge1_scores)),
            'rougeL': float(sum(rougeL_scores) / len(rougeL_scores))
        }

    def process_docx(self, file_path: Path) -> str:
        """Process a DOCX file and return its text content."""
        try:
            doc = Document(str(file_path))  # Convert Path to string
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            return ""

    def process_pdf(self, file_path: Path) -> str:
        """Process a PDF file and return its text content."""
        try:
            with pdfplumber.open(str(file_path)) as pdf:  # Convert Path to string
                return "\n".join([page.extract_text() for page in pdf.pages])
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {str(e)}")
            return ""

def main():
    parser = argparse.ArgumentParser(description='Process medical documents for tagging and summarization')
    parser.add_argument('--input', required=True, help='Input file path (PDF or DOCX)')
    parser.add_argument('--output_dir', required=True, help='Output directory for results')
    args = parser.parse_args()
    
    processor = MedicalTextProcessor(args.input, args.output_dir)
    processor.process_file()
    
    logger.info(f"Processing complete. Results saved in {args.output_dir}")

if __name__ == '__main__':
    main() 